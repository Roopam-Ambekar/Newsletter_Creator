import streamlit as st
import pandas as pd
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.shared import qn
from docx.oxml import OxmlElement

# Helper functions
def add_divider(paragraph):
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_paragraph_format(paragraph):
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    paragraph.paragraph_format.line_spacing = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
        rFonts.set(qn(f"w:{attr}"), "Arial")
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "20")
    rPr.append(sz)

    for tag in ["w:b", "w:i", "w:u"]:
        el = OxmlElement(tag)
        if tag == "w:u":
            el.set(qn("w:val"), "single")
        rPr.append(el)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "4F81BD")
    rPr.append(color)

    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# Main logic function
def process_excel_and_create_word(df, custom_name):
    df_newssheet = df.copy()
    doc = Document()

    # --- Set Moderate Margins ---
    section = doc.sections[0]
    section.top_margin = Pt(36)     
    section.bottom_margin = Pt(36)  
    section.left_margin = Pt(36)    
    section.right_margin = Pt(36)  

    # --- Add Title at the Start ---
    title_para = doc.add_paragraph(custom_name)
    title_para.alignment = 1  # Center alignment
    run = title_para.runs[0]
    run.font.name = "Cambria"
    run.font.size = Pt(20)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

    doc.add_paragraph()  # One-line space

    styles = doc.styles
    if "Supplier Heading" not in styles:
        supplier_style = styles.add_style("Supplier Heading", WD_STYLE_TYPE.PARAGRAPH)
        supplier_style.font.name = "Arial"
        supplier_style.font.size = Pt(12)
        supplier_style.font.bold = True
        supplier_style.font.color.rgb = RGBColor(31, 73, 125)
        supplier_style.paragraph_format.space_before = Pt(12)
        supplier_style.paragraph_format.space_after = Pt(6)

    grouped = df_newssheet.groupby("Supplier")

    first = True

    for supplier, group in grouped:
        if not first:
            doc.add_page_break()
        first = False
        heading = doc.add_paragraph(supplier.upper(), style="Supplier Heading")
        set_paragraph_format(heading)
        set_paragraph_format(doc.add_paragraph())

        for _, row in group.iterrows():
            p1 = doc.add_paragraph(row["Headline"])
            run1 = p1.runs[0]
            run1.font.name = "Calibri"
            run1.font.size = Pt(12)
            run1.bold = True
            run1.font.color.rgb = RGBColor(31, 73, 125)
            p1.paragraph_format.alignment = 3
            set_paragraph_format(p1)

            p2 = doc.add_paragraph(str(row["Date"]))
            r2 = p2.runs[0]
            r2.font.name = "Calibri"
            r2.font.size = Pt(12)
            r2.italic = True
            p2.paragraph_format.alignment = 3
            set_paragraph_format(p2)
            set_paragraph_format(doc.add_paragraph())

            p3 = doc.add_paragraph(str(row["Summary"]))
            r3 = p3.runs[0]
            r3.font.name = "Calibri"
            r3.font.size = Pt(12)
            p3.paragraph_format.alignment = 3
            set_paragraph_format(p3)
            set_paragraph_format(doc.add_paragraph())

            p4 = doc.add_paragraph()
            source_cols = [col for col in df_newssheet.columns if col.startswith("Source")]
            urls = [
                str(row[col]).strip()
                for col in source_cols
                if pd.notna(row[col]) and str(row[col]).strip().lower() != "nan"
            ]
            plural = "links" if len(urls) > 1 else "link"
            run = p4.add_run(
                f"Category: {row['Category']} | Web {plural} to Full Story: "
            )
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(10)

            for i, url in enumerate(urls):
                if i > 0:
                    p4.add_run(", ")
                add_hyperlink(p4, url, "Read More")
            set_paragraph_format(p4)

            divider_para = doc.add_paragraph()
            add_divider(divider_para)
            set_paragraph_format(divider_para)
            set_paragraph_format(doc.add_paragraph())

    # Save to memory
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# --- Streamlit UI ---
st.title("ğŸ“Š ComputerShare Supplier Newsletter Creator")

# ğŸ’¡ Instructions
st.markdown("""
### ğŸ§¾ News Sheet File Format Instructions

Please ensure your Excel file contains the following columns in **this exact order**:

| Column Name | Description |
|-------------|-------------|
| `S. No`     | Serial number (optional, not used in processing) |
| `Supplier`  | Name of the supplier (e.g., Cognizant, Coinbase...) |
| `Category`  | Category of News(e.g. Service Offering, Earnings, Strategic Partnership...) |
| `Shared`    | If the news was sent to the supplier as an alert (Y for Yes/ N for No) |
| `Date`      | Date of the news (e.g., July 2, 2025) |
| `Headline`  | Headline of the news piece|
| `Summary`   | Summary of the news piece |
| `Source 1`  | URL to the full article |
| `Source 2`  | Additional source |

ğŸ“Œ **Note:** Keep column headers exactly as shown, and ensure there's no extra spacing or typos.
""")

st.image("https://github.com/Roopam-Ambekar/Newsletter_Creator/blob/main/Example%20arrangement.jpg?raw=true", caption="ğŸ“„ Example of Correct Excel Format", width = 1000)

# Option selector
option = st.radio(
    "Choose input method:",
    ("ğŸ“‚ Upload Excel File", "ğŸŒ Use Google Sheet")
)

custom_name = st.text_input("Enter name for the Word file (without .docx):", "Enter the Name")

df = None

if option == "ğŸ“‚ Upload Excel File":
    uploaded_file = st.file_uploader("Upload your Excel file", type=["xlsx"])
    if uploaded_file:
        df = pd.read_excel(uploaded_file)

elif option == "ğŸŒ Use Google Sheet":
    st.info("Using data from linked Google Sheet.")
    # Replace with your actual Sheet URL
    sheet_url = "https://docs.google.com/spreadsheets/d/1aErka1fqYFGctlOuXYcjHgfnw9lXMElb7UHqEjh1SWo/export?format=csv"
    try:
        df = pd.read_csv(sheet_url)
    except Exception as e:
        st.error(f"Failed to read Google Sheet: {e}")

if df is not None and custom_name.strip():
    if st.button("ğŸš€ Generate Word Document"):
        word_output = process_excel_and_create_word(df, custom_name)
        st.success("âœ… Word document generated successfully!")

        st.download_button(
            label="ğŸ“¥ Download Word Document",
            data=word_output,
            file_name=f"{custom_name.strip()}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
